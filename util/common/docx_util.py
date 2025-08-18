from pathlib import Path

from docx import Document
from docx.shared import Inches
from loguru import logger

from util.common.constants import Constants


def generate_docx(data: dict, filename: str):
    # 新建 Word 文档
    doc = Document()
    doc.add_heading("视频笔记", level=0)

    # 写 summary
    doc.add_heading("总结", level=1)
    doc.add_paragraph(data.get("summary", ""))

    # 写 highlights
    doc.add_heading("要点", level=1)
    highlights = data.get("highlights", [])
    for item in highlights:
        doc.add_paragraph(f"• {item}")

    # 写 groups 内容
    groups = data.get("groups", [])
    for group in groups:
        title = group.get("title", "")
        summary = group.get("summary", "")
        begin_time = group.get("part_begin_time", "")
        end_time = group.get("part_end_time", "")
        img_url = group.get("img_url", "")
        sentences = group.get("sentences", [])

        doc.add_heading(title, level=1)
        doc.add_paragraph(f"时间：{begin_time} - {end_time}")
        doc.add_paragraph(summary)

        # 插入图片, /upload/img/c786d860-1529-4aac-91b7-5276b8a13677.png
        if img_url:
            try:
                img_url = str(Path(Constants.IMG_DIR) / Path(img_url).name)
                doc.add_picture(img_url, width=Inches(4))
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {img_url}]")

        # 插入句子
        for sentence in sentences:
            begin = sentence.get("begin_time", "")
            end = sentence.get("end_time", "")
            text = sentence.get("text", "")
            doc.add_paragraph(f"{begin}-{end}: {text}")

    # 保存 Word 文档
    save_path = Path(Constants.TMP_DIR) / filename
    doc.save(str(save_path))
    logger.info(f"Word 文档已生成：{save_path}")